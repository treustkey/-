# генератор DOCX документов
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


class DocxGenerator:
    def __init__(self, output_dir="output"):
        # инициализация
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def generate(self, name, documentation_type, system_type, deadline,
                 description, functional_requirements, nonfunctional_requirements=""):
        # генерация документа DOCX

        # создание документа
        doc = Document()

        # настройка полей (по ГОСТ)
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(0.4)
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)

        # заголовок
        title = doc.add_heading('ТЕХНИЧЕСКОЕ ЗАДАНИЕ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # подзаголовок
        subtitle = doc.add_heading(f'{name}', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # ОБЩИЕ СВЕДЕНИЯ
        doc.add_heading('1. ОБЩИЕ СВЕДЕНИЯ', 1)

        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'

        table.cell(0, 0).text = 'Тип документации:'
        table.cell(0, 1).text = documentation_type

        table.cell(1, 0).text = 'Тип системы:'
        table.cell(1, 1).text = system_type or 'Не указан'

        table.cell(2, 0).text = 'Срок выполнения:'
        table.cell(2, 1).text = deadline.strftime('%d.%m.%Y')

        table.cell(3, 0).text = 'Дата создания ТЗ:'
        table.cell(3, 1).text = datetime.now().strftime('%d.%m.%Y')

        doc.add_paragraph()

        # НАЗНАЧЕНИЕ И ЦЕЛИ
        doc.add_heading('2. НАЗНАЧЕНИЕ И ЦЕЛИ СОЗДАНИЯ СИСТЕМЫ', 1)
        doc.add_paragraph(description if description else 'Не указано')

        doc.add_paragraph()

        # ХАРАКТЕРИСТИКА ОБЪЕКТА
        doc.add_heading('3. ХАРАКТЕРИСТИКА ОБЪЕКТА АВТОМАТИЗАЦИИ', 1)
        doc.add_paragraph(f'Объектом автоматизации является {system_type or "система"}.')

        doc.add_paragraph()

        # ТРЕБОВАНИЯ
        doc.add_heading('4. ТРЕБОВАНИЯ К СИСТЕМЕ', 1)

        # функциональные требования
        doc.add_heading('4.1. Функциональные требования', 2)
        if functional_requirements:
            for line in functional_requirements.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip(), style='List Bullet')
        else:
            doc.add_paragraph('Не указаны')

        doc.add_paragraph()

        # нефункциональные требования
        if nonfunctional_requirements:
            doc.add_heading('4.2. Нефункциональные требования', 2)
            for line in nonfunctional_requirements.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip(), style='List Bullet')
            doc.add_paragraph()

        # СОСТАВ И СОДЕРЖАНИЕ РАБОТ
        doc.add_heading('5. СОСТАВ И СОДЕРЖАНИЕ РАБОТ', 1)
        work_stages = [
            'Анализ требований и проектирование системы',
            'Разработка программного обеспечения',
            'Тестирование и отладка',
            'Документирование системы',
            'Внедрение и сопровождение'
        ]
        for stage in work_stages:
            doc.add_paragraph(stage, style='List Number')

        doc.add_paragraph()

        # ПОРЯДОК КОНТРОЛЯ И ПРИЕМКИ
        doc.add_heading('6. ПОРЯДОК КОНТРОЛЯ И ПРИЕМКИ СИСТЕМЫ', 1)
        doc.add_paragraph(
            'Приемка системы осуществляется на основании результатов комплексного '
            'тестирования и соответствия требованиям, изложенным в настоящем техническом задании.'
        )

        # формирование имени файла
        safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).rstrip()
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{safe_name}_{timestamp}.docx"
        output_path = os.path.join(self.output_dir, filename)

        # сохранение
        doc.save(output_path)

        return output_path
